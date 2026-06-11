"""多格式文档解析器。

本文件只负责把本地文件解析成 ParsedDocument。
复杂的入库流程放在 knowledge_service.py，避免解析逻辑和业务编排混在一起。
"""

from __future__ import annotations

import csv
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass
class ParsedDocument:
    """解析器内部使用的文档对象。

    这个对象不对应 MySQL 表，只是把文件解析后的文本传给入库服务。
    """

    title: str
    source_type: str
    source_uri: str
    content: str
    pages: list[dict[str, Any]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentParser:
    """文档解析器。
    文件要先存到本地，然后输入路径进行解析
    支持 pdf、docx、xlsx、csv、md、txt。解析失败时会抛出带中文说明的异常，
    由上层 KnowledgeIngestor 记录到 DocumentIngestResult 中。
    """

    supported_suffixes = {".pdf", ".docx", ".xlsx", ".csv", ".md", ".txt"}

    def parse_file(self, file_path: str | Path) -> ParsedDocument:
        """解析单个文件，并返回统一的文档对象。

        Args:
            file_path: 本地文件路径。

        Returns:
            ParsedDocument: 解析后的文档对象。
        """

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在：{path}")
        if not path.is_file():
            raise ValueError(f"路径不是文件：{path}")

        suffix = path.suffix.lower()
        if suffix not in self.supported_suffixes:
            raise ValueError(f"暂不支持的文件类型：{suffix}")

        if suffix == ".pdf":
            pages = self._parse_pdf(path)
        elif suffix == ".docx":
            pages = self._parse_docx(path)
        elif suffix == ".xlsx":
            pages = self._parse_xlsx(path)
        elif suffix == ".csv":
            pages = self._parse_csv(path)
        elif suffix == ".md":
            pages = self._parse_text_file(path, source_type="md")
        else:
            pages = self._parse_text_file(path, source_type="txt")

        content = "\n\n".join(item["content"] for item in pages if item.get("content"))
        if not content.strip():
            raise ValueError(f"文件没有解析出有效文本：{path}")

        return ParsedDocument(
            title=path.stem,
            source_type=suffix.lstrip("."),
            source_uri=str(path),
            content=content,
            pages=pages,
            metadata={"file_name": path.name, "file_size": path.stat().st_size},
        )

    def _parse_pdf(self, path: Path) -> list[dict[str, Any]]:
        """解析 PDF，每一页作为一个 section。"""

        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ImportError("解析 PDF 需要安装 pypdf") from exc

        reader = PdfReader(str(path))
        pages: list[dict[str, Any]] = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = self._clean_text(text)
            if text:
                pages.append({"page_no": index, "section": f"第 {index} 页", "content": text})
        return pages

    def _parse_docx(self, path: Path) -> list[dict[str, Any]]:
        """解析 Word 文档，正文段落和表格都会尽量转成文本。"""

        try:
            from docx import Document
        except ImportError as exc:
            raise ImportError("解析 Word 文档需要安装 python-docx") from exc

        document = Document(str(path))
        lines: list[str] = []

        for paragraph in document.paragraphs:
            text = self._clean_text(paragraph.text)
            if text:
                lines.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [self._clean_text(cell.text) for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if cells:
                    lines.append(" | ".join(cells))

        content = "\n".join(lines)
        return [{"page_no": None, "section": path.stem, "content": content}]

    def _parse_xlsx(self, path: Path) -> list[dict[str, Any]]:
        """解析 Excel，每个工作表作为一个 section。"""

        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise ImportError("解析 Excel 需要安装 openpyxl") from exc

        workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
        pages: list[dict[str, Any]] = []

        for sheet in workbook.worksheets:
            lines: list[str] = []
            for row in sheet.iter_rows(values_only=True):
                cells = [self._cell_to_text(cell) for cell in row]
                cells = [cell for cell in cells if cell]
                if cells:
                    lines.append(" | ".join(cells))
            content = "\n".join(lines)
            if content:
                pages.append({"page_no": None, "section": sheet.title, "content": content})

        workbook.close()
        return pages

    def _parse_csv(self, path: Path) -> list[dict[str, Any]]:
        """解析 CSV。优先使用 utf-8-sig，失败后尝试 gbk。"""

        rows: list[list[str]] = []
        last_error: UnicodeDecodeError | None = None

        for encoding in ("utf-8-sig", "gbk"):
            try:
                with path.open("r", encoding=encoding, newline="") as file:
                    rows = list(csv.reader(file))
                break
            except UnicodeDecodeError as exc:
                last_error = exc
        else:
            raise UnicodeDecodeError(
                last_error.encoding,
                last_error.object,
                last_error.start,
                last_error.end,
                "CSV 不是 utf-8-sig 或 gbk 编码",
            )

        lines = [" | ".join(self._clean_text(cell) for cell in row if self._clean_text(cell)) for row in rows]
        content = "\n".join(line for line in lines if line)
        return [{"page_no": None, "section": path.stem, "content": content}]

    def _parse_text_file(self, path: Path, source_type: str) -> list[dict[str, Any]]:
        """解析 Markdown 或普通文本。"""

        text = self._read_text_with_fallback(path)
        section = "Markdown 文档" if source_type == "md" else "文本文件"
        return [{"page_no": None, "section": section, "content": self._clean_text(text)}]

    def _read_text_with_fallback(self, path: Path) -> str:
        """读取文本文件，兼容常见中文编码。"""

        last_error: UnicodeDecodeError | None = None
        for encoding in ("utf-8-sig", "utf-8", "gbk"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError as exc:
                last_error = exc

        raise UnicodeDecodeError(
            last_error.encoding,
            last_error.object,
            last_error.start,
            last_error.end,
            "文本不是 utf-8-sig、utf-8 或 gbk 编码",
        )

    def _cell_to_text(self, value: Any) -> str:
        """把表格单元格值转换成干净文本。"""

        if value is None:
            return ""
        return self._clean_text(str(value))

    def _clean_text(self, text: str) -> str:
        """清理文本中的多余空白，保留自然换行。"""

        lines = [" ".join(line.strip().split()) for line in text.splitlines()]
        return "\n".join(line for line in lines if line)
